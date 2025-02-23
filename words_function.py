import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from Math_util import MathUtils
# Function to add borders to a table
class Wordfunctions():
    def add_table_borders(table):
        """
        Add borders to all cells of a table.
        """

        for row in table.rows:  # Iterate over all rows
            for cell in row.cells:  # Iterate over all cells in the row
                cell_properties = cell._element.get_or_add_tcPr()
                borders = OxmlElement('w:tcBorders')

                # Add borders to all sides (top, bottom, left, right)
                for border_name in ['top', 'bottom', 'left', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')  # Border style: single line
                    border.set(qn('w:sz'), '4')  # Border size: 4 (1/8th of a point)
                    border.set(qn('w:space'), '0')  # No spacing
                    border.set(qn('w:color'), '000000')  # Border color: black
                    borders.append(border)

                cell_properties.append(borders)

    # Function to insert a page break before a table and adjust the table size
    def insert_page_break_and_adjust_table(doc, table):
        """
        Adjust the table to fit within one page (optional resizing) and prepare space for the page break.
        """
        # Resize the table to fit within the page (you can adjust the width here)
        for row in table.rows:
            for cell in row.cells:
                cell.width = Pt(100)  # Adjust this value as needed for your table
    
        # Optionally, set autofit for the table to ensure it resizes to fit within the page width
        table.autofit = True

# Function to detect values inside the table (excluding first row and column) and calculate percentage

class Wordsfunctions2:
    def __init__(self, filename=None):
        """Initialize the WordHandler class. If a filename is provided, open that document."""
        self.cwd = os.getcwd()  # Get current working directory
        if filename and os.path.exists(os.path.join(self.cwd, f"{filename}.docx")):
            self.doc = Document(os.path.join(self.cwd, f"{filename}.docx"))
            self.filename = filename
        else:
            self.doc = Document()
            self.filename = None

    def write_statement(self, statement):
        """Write a statement (paragraph) to the Word document."""
        self.doc.add_paragraph(statement)

    def save_document(self, filename):
        """Save the document with the given filename."""
        self.filename = filename
        self.doc.save(os.path.join(self.cwd, f"{filename}.docx"))

    def read_statements(self, filename):
        """Read all paragraphs from a Word document while ignoring tables."""
        file_path = os.path.join(self.cwd, f"{filename}.docx")
        if not os.path.exists(file_path):
            return ["Error: File does not exist."]
        
        doc = Document(file_path)  # Load the document
        return [para.text for para in doc.paragraphs if not para._element.getparent().tag.endswith('tbl')]

    def delete_document(self, filename):
        """Delete the specified Word document if it exists."""
        file_path = os.path.join(self.cwd, f"{filename}.docx")
        if os.path.exists(file_path):
            os.remove(file_path)
            return f"File '{filename}.docx' deleted successfully."
        else:
            return "Error: File does not exist."

    def list_documents(self):
        """List all .docx files in the current working directory."""
        return [f for f in os.listdir(self.cwd) if f.endswith(".docx")]

    def get_working_directory(self):
        """Return the current working directory."""
        return self.cwd
