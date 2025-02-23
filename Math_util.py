from collections import defaultdict
from collections import Counter

from docx.shared import RGBColor

import re


class MathUtils:

    def extract_and_combine_percentages(self,lists_of_values):
        """
    Combines percentages for each value (e.g., 'Value X') across multiple lists,
    summing the percentages if the same value appears multiple times.
    """
        combined_percentages = defaultdict(float)

        for lst in lists_of_values:
            for item in lst:
                match = re.search(r'Value (\d+): \((\d+(\.\d+)?)%\)', item)
                if match:
                    value_name = f"Value {match.group(1)}"
                    percentage_value = float(match.group(2))
                    combined_percentages[value_name] += percentage_value
                else:
                    print(f"Warning: Could not extract percentage from '{item}'")

        return sorted(combined_percentages.items())  # Return sorted values


        
    def calculate_percentage_and_add_paragraph(doc, table):
        """
    Calculate the percentage of values in a table (excluding first row and column)
    and add a paragraph under the table. The percentage calculation is based on the frequency of each value.
    The highest percentage is highlighted with different colors based on occurrences:
    - 3 times: Dark red
    - 4 times: Dark dark red
    - 5 or more times: Pink
    
    (This function work for words with scri[:short_probability_scanner.py])
        """
    # Collect all the numeric values from the table (excluding first row and first column)
        table_values = []
        total_value = 0

        # Gather all values (excluding first row and first column)
        for row_idx, row in enumerate(table.rows[1:]):  # Exclude the first row (header)
            for col_idx, cell in enumerate(row.cells[1:]):  # Exclude the first column
                try:
                    value = float(cell.text)
                    table_values.append(value)
                    total_value += 1  # Count total occurrences of values
                except ValueError:
                    continue  # Skip if the cell value is not a number

    # If no valid numerical values, inform the user and return
        if total_value == 0:
            doc.add_paragraph("This table contains no numerical values to calculate percentages.")
            return

    # Count the frequency of each value in the table
        value_counts = Counter(table_values)
     # Sort the values by frequency in descending order
        sorted_values = sorted(value_counts.items(), key=lambda x: x[1], reverse=True)

    # Add the results to the paragraph
        for value, count in sorted_values:
            percentage = (count / total_value) * 100
            value_text = f"Value {value:.2f} appears {count} times, which is {percentage:.2f}% of the total."
        # Determine the color based on the count
            if count == 2:
                color = RGBColor(255, 0, 0)  # Red
            elif count == 3:
                color = RGBColor(139, 0, 0)  # Dark red
            elif count == 4:
                color = RGBColor(139, 0, 0)  # Dark dark red
            elif count >= 5:
                color = RGBColor(255, 182, 193)  # Pink
            else:
                color = None  # Default color (black)

            # Apply the color to the text
            p = doc.add_paragraph()
            run = p.add_run(value_text)

            if color:
                run.font.color.rgb = color
        
        print("Added summary paragraph with percentage calculations.")