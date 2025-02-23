import os
import re
from collections import Counter
from docx import Document
from docx.shared import RGBColor
from Data_storage_Lib import total_percent_list  # Import the external percentage list

class summary_percentage_writer:
    def __init__(self, input_filename="output_tables.docx", output_filename="percentage_total_percent.docx"):
        """
        Initialize the writer with an input and output filename.
        A new summary document is created, and external percentages are parsed.
        """
        self.input_filename = input_filename
        self.output_filename = output_filename
        self.summary_doc = Document()
        self.external_percentages = self.parse_total_percent_list()

    def extract_percentage(self, value_string):
        """
        Extract the numerical value and percentage from a string like 'Value 22: (2.38%)'.
        Returns a tuple (value, percentage) or (None, None) if the pattern doesn't match.
        """
        match = re.match(r"Value (\d+): \(([\d.]+)%\)", value_string)
        if match:
            value = float(match.group(1))
            percentage = float(match.group(2))
            return value, percentage
        return None, None

    def parse_total_percent_list(self):
        """
        Parse total_percent_list into a dictionary {value: percentage}.
        For example: {22.0: 2.38, 31.0: 2.32, ...}
        """
        parsed_data = {}
        for item in total_percent_list:
            value, percentage = self.extract_percentage(item)
            if value is not None:
                parsed_data[value] = percentage
        return parsed_data

    def calculate_percentage_and_add_paragraph(self, doc, table):
        """
        Calculate the percentage of numeric values in a table (excluding the first row and column)
        and add summary paragraphs to the passed document (doc).
        The summary is updated using external percentages from the parsed total_percent_list.
        Additionally, a paragraph is added that lists values from 1 to 49 that did not occur,
        and this missing values paragraph is written in yellow.
        """
        # Collect all numeric values (skip header row and first column)
        table_values = []
        total_value = 0
        for row in table.rows[1:]:
            for cell in row.cells[1:]:
                try:
                    value = float(cell.text)
                    table_values.append(value)
                    total_value += 1
                except ValueError:
                    continue

        if total_value == 0:
            doc.add_paragraph("This table contains no numerical values to calculate percentages.")
            return

        # Count the frequency of each value and sort by frequency in descending order
        value_counts = Counter(table_values)
        sorted_values = sorted(value_counts.items(), key=lambda x: x[1], reverse=True)

        # Add a summary paragraph for each unique value found in the table
        for value, count in sorted_values:
            percentage = (count / total_value) * 100
            external_percentage = self.external_percentages.get(value, 0)
            total_percentage = percentage * external_percentage  # Using multiplication as in your code
            value_text = f"Value {value:.2f} appears {count} times, which is {percentage:.2f}% of the total."
            total_text = f"Updated percentage (including external data): {total_percentage:.2f}%."

            # Determine the font color based on the count
            if count == 2:
                color = RGBColor(255, 0, 0)  # Red
            elif count == 3:
                color = RGBColor(139, 0, 0)  # Dark red
            elif count == 4:
                color = RGBColor(139, 0, 0)  # Dark dark red
            elif count >= 5:
                color = RGBColor(255, 182, 193)  # Pink
            else:
                color = None  # Default (black)

            # Add the paragraph with the calculated value
            p = doc.add_paragraph()
            run = p.add_run(value_text)
            if color:
                run.font.color.rgb = color

            # Add the updated percentage on a new paragraph
            doc.add_paragraph(total_text)

        # Now compute and add a paragraph listing values from 1 to 49 that did NOT occur in the table.
        found_values = set()
        # Convert found values to integers (only if they fall between 1 and 49)
        for val in table_values:
            try:
                val_int = int(round(val))
                if 1 <= val_int <= 49:
                    found_values.add(val_int)
            except:
                continue
        missing_values = [i for i in range(1, 50) if i not in found_values]
        missing_text = f"Values from 1 to 49 that have (0%): {missing_values}"
        # Add the missing values paragraph in yellow
        p_missing = doc.add_paragraph()
        run_missing = p_missing.add_run(missing_text)
        run_missing.font.color.rgb = RGBColor(153, 120, 0) # Yellow color

    def process_document(self):
        """
        Process the input document: loop through all tables, add a header and page break before
        each table's summary, and then call the function to calculate percentages.
        Finally, save the summary document.
        """
        if not os.path.exists(self.input_filename):
            print(f"Error: {self.input_filename} not found!")
            return

        # Open the input document
        doc = Document(self.input_filename)
        if not doc.tables:
            print("No tables found in the document.")
            return

        # Process each table
        for index, table in enumerate(doc.tables):
            print(f"Processing Table {index + 1}...")
            # Add a page break for all tables except the first
            if index > 0:
                self.summary_doc.add_page_break()
            # Add a header for the table summary
            self.summary_doc.add_paragraph(f"Table {index + 1}", style="Heading 1")
            # Append table summary to the summary document
            self.calculate_percentage_and_add_paragraph(self.summary_doc, table)

        # Save the summary document
        self.summary_doc.save(self.output_filename)
        print(f"✅ All tables processed successfully! Summary saved as '{self.output_filename}'.")

    def run(self):
        """Run the entire process."""
        self.process_document()

