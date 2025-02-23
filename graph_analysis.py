import sys
import os
import numpy as np
import pyqtgraph as pg
import importlib.util
import datetime
import docx  # pip install python-docx

from PySide6.QtWidgets import QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QPushButton, QTextEdit
from PySide6.QtCore import QTimer, Qt
from PySide6.QtGui import QFont
from sklearn.linear_model import LinearRegression
from sklearn.neighbors import KNeighborsRegressor
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import mean_squared_error
import pyqtgraph.exporters

class GraphAnalysisWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.initUI()
        # Load data from Data_storage_Lib.py
        ds_path = os.path.join(os.getcwd(), "Data_storage_Lib.py")
        spec = importlib.util.spec_from_file_location("Data_storage_Lib", ds_path)
        ds_module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(ds_module)
        self.graph_plot_info = getattr(ds_module, "graph_plot_info", {})
        self.predict_dict = getattr(ds_module, "predict_dict", {})
        self.prediction_index = 1
        self.is_total_predict = False
        self.best_prediction_key = None
        self.flash_timer = QTimer(self)
        self.flash_timer.timeout.connect(self.toggle_flash_color)
        self.flash_state = False
        if self.predict_dict:
            self.show_prediction(self.prediction_index)
        else:
            self.text_edit.setPlainText("No prediction data found in Data_storage_Lib.py.")

    def initUI(self):
        layout = QHBoxLayout(self)
        # Left: pyqtgraph plot widget
        self.plot_widget = pg.PlotWidget(title="Real vs Virtual Line Plot")
        self.plot_widget.setLabel('left', 'Y', units='(percentage)')
        self.plot_widget.setLabel('bottom', 'X', units='(point number)')
        self.plot_widget.showGrid(x=True, y=True)
        layout.addWidget(self.plot_widget, stretch=3)
        # Right: Information panel and buttons
        right_layout = QVBoxLayout()
        layout.addLayout(right_layout, stretch=1)
        self.text_edit = QTextEdit()
        self.text_edit.setReadOnly(True)
        self.text_edit.setStyleSheet("background-color: black; color: white;")
        right_layout.addWidget(self.text_edit)
        btn_layout = QHBoxLayout()
        self.next_button = QPushButton("Next Prediction")
        self.next_button.clicked.connect(self.on_next_button_clicked)
        btn_layout.addWidget(self.next_button)
        self.total_predict_button = QPushButton("Total Predict")
        self.total_predict_button.clicked.connect(self.on_total_predict_button_clicked)
        btn_layout.addWidget(self.total_predict_button)
        self.save_graph_button = QPushButton("Save Graph Report")
        self.save_graph_button.clicked.connect(self.on_save_graph_clicked)
        btn_layout.addWidget(self.save_graph_button)
        right_layout.addLayout(btn_layout)

    def on_next_button_clicked(self):
        self.prediction_index += 1
        if self.prediction_index > len(self.predict_dict):
            self.prediction_index = 1
        self.plot_widget.clear()
        self.flash_timer.stop()
        self.show_prediction(self.prediction_index)

    def on_total_predict_button_clicked(self):
        self.is_total_predict = True
        self.plot_widget.clear()
        self.text_edit.clear()
        self.plot_all_predictions()
        self.display_total_predict_info()
        self.plot_green_line()
        self.plot_widget.setInteractive(True)

    def on_save_graph_clicked(self):
        try:
            exporter = pg.exporters.ImageExporter(self.plot_widget.plotItem)
            png_filename = "graph_export.png"
            exporter.export(png_filename)
            doc = docx.Document()
            doc.add_heading("Graph Analysis Report", level=1)
            doc.add_paragraph("Generated on: " + datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
            doc.add_heading("Graph Information", level=2)
            info_text = self.text_edit.toPlainText()
            doc.add_paragraph(info_text)
            doc.add_heading("Graph Image", level=2)
            doc.add_picture(png_filename, width=docx.shared.Inches(6))
            report_filename = "graph_analysis_report.docx"
            doc.save(report_filename)
            self.text_edit.append("\nReport saved as " + report_filename)
        except Exception as e:
            self.text_edit.append("\nError saving report: " + str(e))

    def show_prediction(self, i_info):
        self.plot_widget.clear()
        prediction_key = f'predict_{i_info}'
        real_line = self.predict_dict.get(prediction_key, [])
        if not real_line:
            self.text_edit.setPlainText(f"No data for {prediction_key}")
            return
        virtual_line = self.generate_virtual_line(real_line)
        yellow_virtual_line, yellow_date = self.generate_yellow_virtual_line(real_line)
        self.plot_data(real_line, virtual_line, yellow_virtual_line, prediction_key, yellow_date)
        self.update_message_display(prediction_key, yellow_date, real_line, virtual_line, yellow_virtual_line)

    def plot_data(self, real_line, virtual_line, yellow_virtual_line, prediction_key, yellow_date):
        x_real, y_real = zip(*real_line)
        self.plot_widget.plot(x_real, y_real, pen=pg.mkPen('b', width=2), symbol='o', symbolBrush='b')
        x_virtual, y_virtual = zip(*virtual_line)
        red_pen = pg.mkPen(color='r', width=2, style=Qt.PenStyle.DashLine)
        self.plot_widget.plot(x_virtual, y_virtual, pen=red_pen)
        x_yellow, y_yellow = zip(*yellow_virtual_line)
        yellow_pen = pg.mkPen(color='y', width=2, style=Qt.PenStyle.DashLine)
        self.plot_widget.plot(x_yellow, y_yellow, pen=yellow_pen)
        self.adjust_range(x_real, y_real)

    def adjust_range(self, x_data, y_data):
        if self.is_total_predict:
            return
        y_min, y_max = min(y_data), max(y_data)
        y_buffer = 0.03 * (y_max - y_min)
        self.plot_widget.setYRange(y_min - y_buffer, y_max + y_buffer)

    def plot_green_line(self):
        all_y_values = []
        for key in self.predict_dict:
            real_line = self.predict_dict[key]
            if real_line:
                _, y_real = zip(*real_line)
                all_y_values.extend(y_real)
        if all_y_values:
            highest_y_value = max(all_y_values)
            self.plot_widget.plot([25, 25], [0, highest_y_value * 1.03],
                                  pen=pg.mkPen('g', width=2))

    def update_message_display(self, prediction_key, yellow_date, real_line, virtual_line, yellow_virtual_line):
        x_real, y_real = zip(*real_line)
        x_virtual, y_virtual = zip(*virtual_line)
        x_yellow, y_yellow = zip(*yellow_virtual_line)
        mse_real_vs_virtual = mean_squared_error(y_real, y_virtual)
        mse_real_vs_yellow = mean_squared_error(y_real, y_yellow)
        x_real_arr = np.array(x_real).reshape(-1, 1)
        lr = LinearRegression()
        lr.fit(x_real_arr, y_real)
        slope = lr.coef_[0]
        intercept = lr.intercept_
        r2 = lr.score(x_real_arr, y_real)
        message = (
            f"Graph Information:\n\n"
            f"Prediction Key: {prediction_key}\n"
            f"Closest Trend Date: {yellow_date}\n\n"
            f"Advanced Math Details:\n"
            f"  - Slope: {slope:.4f}\n"
            f"  - Intercept: {intercept:.4f}\n"
            f"  - R² Score: {r2:.4f}\n\n"
            f"Error Metrics:\n"
            f"  - MSE (Blue vs Red): {mse_real_vs_virtual:.4f}\n"
            f"  - MSE (Blue vs Yellow): {mse_real_vs_yellow:.4f}\n\n"
            f"Explanation:\n"
            f"  - Blue line: Actual prediction data.\n"
            f"  - Red dashed line: Linear regression fit.\n"
            f"  - Yellow dashed line: Closest historical trend (via KNN).\n"
            f"  - Higher R² and lower MSE indicate a better fit.\n"
        )
        self.text_edit.setPlainText(message)

    def plot_all_predictions(self):
        self.plot_widget.clear()
        for i_info in range(1, len(self.predict_dict) + 1):
            prediction_key = f'predict_{i_info}'
            real_line = self.predict_dict.get(prediction_key, [])
            if not real_line:
                continue
            virtual_line = self.generate_virtual_line(real_line)
            yellow_virtual_line, yellow_date = self.generate_yellow_virtual_line(real_line)
            self.plot_data(real_line, virtual_line, yellow_virtual_line, prediction_key, yellow_date)

    def display_total_predict_info(self):
        total_info = "Total Predictions Analysis:\n\n"
        total_mse = []
        for i_info in range(1, len(self.predict_dict) + 1):
            prediction_key = f'predict_{i_info}'
            real_line = self.predict_d
