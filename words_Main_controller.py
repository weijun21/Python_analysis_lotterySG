import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from Math_util import MathUtils
from words_function import Wordfunctions
# Function to generate the Word document
class ExcelToWordConverter: 
    def short_prob_scanner(self,filename, sheet, start_col, end_col, row_set):
        """
        Generate a Word document from the Excel sheet and insert a page break after each table.
    
    Arguments:
    - filename: The path to the Excel file.
    - sheet: The name of the sheet to read.
    - start_col: The starting column (inclusive).
    - end_col: The ending column (inclusive).
    - row_set: Number of rows to include in each table.
        """
    # Get the current working directory
        cwd = os.getcwd()
        print("Current Working Directory:", cwd)

    # Check if the file exists
        if not os.path.exists(filename):
            print(f"Error: File '{filename}' does not exist.")
            return  # Exit the function if the file is not found

    # Read the Excel sheet into a DataFrame
        print(f"Reading sheet '{sheet}' from Excel file...")
        df = pd.read_excel(filename, sheet_name=sheet, header=0, usecols=f"{start_col}:{end_col}")
        print(f"Successfully read {len(df)} rows and {len(df.columns)} columns.")

    # Create a Word document
        doc = Document()

    # Iterate over the DataFrame in chunks of `row_set` rows
        print(f"Processing data in chunks of {row_set} rows...")
        for i in range(0, len(df), row_set):
            # Extract `row_set` rows
            chunk = df.iloc[i:i+row_set, :]
            print(f"\nProcessing rows {i+1} to {min(i+row_set, len(df))}...")

            # Add a heading for the table
            table_number = i // row_set + 1
            doc.add_heading(f'Table {table_number}', level=1)
            print(f"Added heading for Table {table_number}.")

            # Create a table in Word with the correct number of rows and columns
            num_rows = len(chunk) + 1  # 1 header row + number of data rows
            num_cols = len(df.columns)  # Number of columns (B:H)
            print(f"Creating table with {num_rows} rows and {num_cols} columns...")

            table = doc.add_table(rows=num_rows, cols=num_cols)

        # Add the header row
            for j, header in enumerate(df.columns):
                table.cell(0, j).text = header
            print("Header row added.")

            # Add the data rows
            for row_idx, row in chunk.iterrows():
                for col_idx, value in enumerate(row):
                    # Calculate the correct row index in the table
                    table_row_idx = (row_idx % row_set) + 1
                    table.cell(table_row_idx, col_idx).text = str(value)
            print(f"Added {len(chunk)} data rows to the table.")

            # Add borders to the table
            Wordfunctions.add_table_borders(table)
            print("Added borders to the table.")

            # Insert table and ensure no page break before each table
            Wordfunctions.insert_page_break_and_adjust_table(doc, table)
            print("Adjusted table for page size.")

            # Perform table detection and percentage calculation
        
            MathUtils.calculate_percentage_and_add_paragraph(doc, table)
            print("Calculated percentage and added paragraph.")

            # Insert a page break after the table's content, including the paragraph with percentage calculations
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            print("Inserted page break after the table and statement.")

        # Save the Word document
        output_file = os.path.join(cwd, 'output_tables.docx')
        doc.save(output_file)
        print(f"\nWord document saved successfully at: {output_file}")
    


