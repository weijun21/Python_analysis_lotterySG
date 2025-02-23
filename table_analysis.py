import re
from collections import defaultdict
from docx import Document
import pandas as pd

class TablesAnalysis:
    def __init__(self, input_tables_filename="output_tables.docx",
                 summary_filename="percentage_total_percent.docx",
                 output_filename="table_analysis.docx"):
        """
        Initialize with filenames:
          - input_tables_filename: document containing original tables.
          - summary_filename: document containing percentage mapping.
          - output_filename: document to store analyzed tables.
        """
        self.input_filename = input_tables_filename
        self.summary_filename = summary_filename
        self.output_filename = output_filename

        self.input_doc = Document(self.input_filename)
        self.summary_doc = Document(self.summary_filename)
        self.analysis_doc = Document()

        # Extract table-specific percentage mapping
        self.table_percentage_mapping = self.extract_table_specific_percentage_mapping()

    def extract_table_specific_percentage_mapping(self):
        """
        Extracts percentage mappings per table from summary document.
        """
        table_mappings = defaultdict(list)
        current_table_index = 0
        paragraphs = self.summary_doc.paragraphs

        for i in range(len(paragraphs) - 1):
            text = paragraphs[i].text.strip()
            next_text = paragraphs[i + 1].text.strip()

            # Detecting Table Headers
            if text.startswith("Table "):
                current_table_index += 1  # Move to next table

            if text.startswith("Value"):
                match_val = re.search(r"Value\s+([\d\.]+)", text)
                match_perc = re.search(r"which is\s+([\d\.]+)% of the total", text)
                match_toperc = re.search(r"Updated percentage \(including external data\): ([\d\.]+)%", next_text)

                if match_val and match_perc and match_toperc:
                    try:
                        value = int(round(float(match_val.group(1))))
                        original_percentage = float(match_perc.group(1))
                        updated_percentage = float(match_toperc.group(1))

                        # Store per table index
                        table_mappings[current_table_index].append((value, original_percentage, updated_percentage))

                    except ValueError:
                        continue

        return dict(table_mappings)

    def extract_comparable_values(self, table):
        """
        Extracts all numeric values from a table for comparison, 
        EXCLUDING first row and first column.
        """
        values = set()
        for i, row in enumerate(table.rows[1:], start=1):  # Ignore header row (Row 1)
            for j, cell in enumerate(row.cells[1:], start=1):  # Ignore first column (Column 1)
                try:
                    value = int(round(float(cell.text.strip())))
                    values.add(value)
                except ValueError:
                    continue
        return values

    def process_tables(self):
        """
        Processes each table, ensuring each table compares with the next one.
        """
        tables = self.input_doc.tables
        num_tables = len(tables)

        # Extract and store data from all tables first (excluding row 1 & column 1 for comparison)
        extracted_comparable_values = [self.extract_comparable_values(table) for table in tables]

        for idx, table in enumerate(tables):
            print(f"Processing Table {idx + 1}...")

            num_rows = len(table.rows)
            num_cols = len(table.rows[0].cells) if num_rows > 0 else 0

            # Add table header
            self.analysis_doc.add_paragraph(f"Table {idx + 1}", style="Heading 1")
            analysis_table = self.analysis_doc.add_table(rows=num_rows, cols=num_cols)
            analysis_table.style = table.style

            # Get the current table's comparable values
            current_table_values = extracted_comparable_values[idx]

            # Compare with next table's values
            if idx + 1 < num_tables:
                next_table_values = extracted_comparable_values[idx + 1]  # Compare with the next table
            else:
                next_table_values = set()  # If last table, no next table

            list_pattern_common = []
            list_recentpercent_common = []

            for i, row in enumerate(table.rows):  # Keep all rows
                row_pattern = []
                row_recentpercent = []
                
                for j, cell in enumerate(row.cells):  # Keep all columns
                    original_text = cell.text.strip()
                    new_text = original_text
                    
                    # **Only compare values from column 2 onward and row 2 downward**
                    if i > 0 and j > 0:
                        try:
                            value = int(round(float(original_text)))
                            
                            # Get percentages specific to this table
                            percentage, match_percentage = 0.0, 0.0
                            if idx + 1 in self.table_percentage_mapping:
                                for val, orig_perc, upd_perc in self.table_percentage_mapping[idx + 1]:
                                    if val == value:
                                        percentage = orig_perc
                                        match_percentage = upd_perc
                                        break

                            # Correctly detect missing values
                            if value not in next_table_values:
                                recent_percentage = 100.0  # Value is missing in the next table
                            else:
                                recent_percentage = 0.0  # Value still exists in the next table

                            new_text = f"{original_text}\n({percentage:.2f}%, {recent_percentage}%, {match_percentage:.2f}%)"
                            row_pattern.append(f"{percentage:.2f}%")
                            row_recentpercent.append(f"{recent_percentage}%")

                        except ValueError:
                            new_text = original_text

                    analysis_table.cell(i, j).text = new_text

                if row_pattern:
                    list_pattern_common.append(row_pattern)
                if row_recentpercent:
                    list_recentpercent_common.append(row_recentpercent)

            # Append findings
            conclusion_text = (
                f"list_pattern_common:\n" + "\n".join(
                    [f"Row {i}: {pattern}" for i, pattern in enumerate(list_pattern_common, start=1)]
                ) + "\n\n"
                f"list_recentpercent_common:\n" + "\n".join(
                    [f"Row {i}: {recent}" for i, recent in enumerate(list_recentpercent_common, start=1)]
                )
            )
            self.analysis_doc.add_paragraph(conclusion_text)
            self.analysis_doc.add_page_break()

        print(f"Finished processing all tables.")

    def save_analysis(self):
        """
        Saves the processed analysis document.
        """
        self.analysis_doc.save(self.output_filename)
        print(f"✅ Analysis saved as '{self.output_filename}'.")

    def run(self):
        """
        Runs the full analysis process.
        """
        self.process_tables()
        self.save_analysis()

