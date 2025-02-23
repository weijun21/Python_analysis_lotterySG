import os
import re
import random
import importlib.util
import sys
from datetime import datetime
from docx import Document

def load_data_storage(file_path):
    spec = importlib.util.spec_from_file_location("Data_storage_Lib", file_path)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module

def parse_table_analysis(file_path):
    doc = Document(file_path)
    trends = []
    list_pattern_common = []
    for table in doc.tables:
        for row in table.rows[1:]:  # Skip header row
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) > 0 and re.match(r"\d{4}-\d{2}-\d{2}", cells[0]):
                date = cells[0]
                numbers = []
                percentages = []
                for cell in cells[1:8]:
                    match = re.search(r"(\d+)", cell)
                    percent_match = re.search(r"([\d.]+)%", cell)
                    if match:
                        numbers.append(int(match.group()))
                    if percent_match:
                        percentages.append(float(percent_match.group().replace('%', '')))
                if len(numbers) == 7 and len(percentages) == 7:
                    trends.append((date, tuple(numbers), tuple(percentages)))
                    list_pattern_common.append((date, tuple(numbers), tuple(percentages)))
    return trends, list_pattern_common

def parse_percentage_total(file_path):
    doc = Document(file_path)
    overall = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        match = re.search(r"Value\s+(\d+(?:\.\d+)?)\s+appears.*Updated percentage.*:\s*([\d\.]+)%", text)
        if match:
            try:
                candidate = int(float(match.group(1)))
                percent = float(match.group(2))
                overall[candidate] = percent
            except Exception:
                continue
    return overall

def select_trend(trends, list_pattern_common, mode):
    if not trends or not list_pattern_common:
        return None, None, "No historical trends found."
    mode_lower = mode.lower()
    if mode_lower == "random":
        selected_trend = random.choice(trends)
        selected_pattern = random.choice(list_pattern_common)
        return selected_trend, selected_pattern, f"Randomly selected trend: {selected_trend}"
    elif mode_lower == "top3":
        trend_counts = {}
        for trend in trends:
            trend_counts[trend] = trend_counts.get(trend, 0) + 1
        sorted_trends = sorted(trend_counts.items(), key=lambda x: x[1], reverse=True)[:3]
        selected_trend = random.choice([t[0] for t in sorted_trends])
        selected_pattern = random.choice(list_pattern_common)
        return selected_trend, selected_pattern, f"Selected one of the Top 3 Trends: {selected_trend}"
    elif mode_lower == "single most frequency":
        trend_counts = {}
        for trend in trends:
            trend_counts[trend] = trend_counts.get(trend, 0) + 1
        selected_trend = max(trend_counts, key=trend_counts.get)
        selected_pattern = random.choice(list_pattern_common)
        return selected_trend, selected_pattern, f"Applied most frequent trend: {selected_trend}"
    return None, None, "Invalid trend selection."

def generate_predictions(num_results, trend_mode):
    print(f"DEBUG: Starting prediction generation - Mode: {trend_mode}, Count: {num_results}")
    
    cwd = os.getcwd()
    ds_file = os.path.join(cwd, "Data_storage_Lib.py")
    table_file = os.path.join(cwd, "table_analysis.docx")
    pct_file = os.path.join(cwd, "percentage_total_percent.docx")
    result_file = os.path.join(cwd, "combination_analysis_result.txt")

    ds_module = load_data_storage(ds_file)
    overall = parse_percentage_total(pct_file)
    trends, list_pattern_common = parse_table_analysis(table_file)

    results = []
    predict_dict = {}

    for i in range(num_results):
        print(f"DEBUG: Generating prediction {i+1}/{num_results}")
        trend, pattern_data, trend_info = select_trend(trends, list_pattern_common, trend_mode)
        if pattern_data:
            pattern_date, pattern_numbers, pattern_percentages = pattern_data
        else:
            pattern_date, pattern_numbers, pattern_percentages = ("N/A", [], [])
        numbers = list(range(1, 50))
        random.shuffle(numbers)
        selected_numbers = sorted(numbers[:7])
        paired = [(x, y) for x, y in zip(selected_numbers, pattern_percentages)]
        predict_dict[f"predict_{i + 1}"] = paired

        prediction_text = (
            f"🎯 Final Combination: {' '.join(map(str, selected_numbers))}\n"
            f"📊 {trend_info}\n"
            f"📅 Applied list_pattern_common Date: {pattern_date}\n"
            f"📋 Applied list_pattern_common Numbers: {pattern_numbers}\n"
            f"📊 Applied list_pattern_common Percentage: {pattern_percentages}\n"
            f"predict_{i + 1}=[{', '.join(f'({x},{y})' for x, y in paired)}]\n"
        )
        results.append(prediction_text)

    result_text = "\n".join(results)
    print(f"DEBUG: Writing results to {result_file}")
    with open(result_file, "w", encoding="utf-8") as f:
        f.write(result_text)

    # Only update Data_storage_Lib.py if the prediction count is low enough
    if num_results <= 20:
        print("DEBUG: Updating Data_storage_Lib.py")
        with open(ds_file, "r", encoding="utf-8") as f:
            lines = f.readlines()
        new_lines = [line for line in lines if "predict_dict" not in line]
        new_lines.append("\n# Updated predictions dictionary\n")
        new_lines.append("predict_dict = " + repr(predict_dict) + "\n")
        with open(ds_file, "w", encoding="utf-8") as f:
            f.write("".join(new_lines))
    else:
        print("DEBUG: Skipping update of Data_storage_Lib.py for high prediction count")

    print("DEBUG: Finished prediction generation")
    return result_text

if __name__ == "__main__":
    if len(sys.argv) >= 3:
        trend_mode = sys.argv[1]
        try:
            prediction_count = int(sys.argv[2])
        except ValueError:
            print("ERROR: Invalid prediction count")
            sys.exit(1)
    else:
        print("Main: Choose Trend Mode (Random/Top3/Single Most Frequency mode)")
        trend_mode = input().strip()
        print("Main: Choose How Many Predictions?")
        try:
            prediction_count = int(input().strip())
        except ValueError:
            print("ERROR: Invalid prediction count")
            sys.exit(1)

    generate_predictions(prediction_count, trend_mode)
    print("Main: Combination Analysis Completed.")
